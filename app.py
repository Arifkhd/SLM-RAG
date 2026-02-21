import os
import re
import sys
import pickle
import shutil

from flask import (
    Flask, request, jsonify, render_template,
    redirect, url_for, session, flash
)
from werkzeug.utils import secure_filename
from dotenv import load_dotenv
from supabase import create_client
from langchain_groq import ChatGroq


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

load_dotenv(dotenv_path=os.path.join(BASE_DIR, ".env"))

HF_TOKEN = os.getenv("HF_TOKEN") or os.getenv("HUGGINGFACEHUB_API_TOKEN")
if HF_TOKEN:
    os.environ["HF_TOKEN"] = HF_TOKEN

from src.vectorstore import FaissVectorStore as QAFaissVectorStore
from src.data_loader import load_all_documents as load_qa_documents
from src_code.vectorstore import FaissVectorStore as SummFaissVectorStore


PERSIST_DIR = os.getenv("PERSIST_DIR", os.path.join(BASE_DIR, "faiss_store"))
SUMMARY_DIR = os.getenv("SUMMARY_DIR", os.path.join(BASE_DIR, "faiss_store_2"))

EMBEDDING_MODEL = os.getenv("EMBEDDING_MODEL", "all-MiniLM-L6-v2")
LLM_MODEL = os.getenv("LLM_MODEL", "llama-3.1-8b-instant")

TOP_K_DEFAULT = int(os.getenv("TOP_K", "3"))
TOP_K_SUMMARY = int(os.getenv("TOP_K_SUMMARY", "5"))

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY = os.getenv("SUPABASE_ANON_KEY")
FLASK_SECRET_KEY = os.getenv("FLASK_SECRET_KEY", "change-me")

if not SUPABASE_URL or not SUPABASE_ANON_KEY:
    raise RuntimeError("Missing SUPABASE_URL or SUPABASE_ANON_KEY in .env")
if not GROQ_API_KEY:
    raise RuntimeError("Missing GROQ_API_KEY in .env")

ALLOWED_EXT = {".pdf", ".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff", ".txt", ".docx"}
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

QA_INDEX_NAME = "faiss.index"
QA_META_NAME = "metadata.pkl"

SUMM_INDEX_NAME = "faiss_summarization.index"
SUMM_META_NAME = "metadata_summarization.pkl"
SUMMARY_PKL_PATH = os.path.join(SUMMARY_DIR, SUMM_META_NAME)


app = Flask(__name__)
app.secret_key = FLASK_SECRET_KEY

supabase = create_client(SUPABASE_URL, SUPABASE_ANON_KEY)
llm = ChatGroq(groq_api_key=GROQ_API_KEY, model_name=LLM_MODEL)

def allowed_file(filename: str) -> bool:
    ext = os.path.splitext(filename)[1].lower()
    return ext in ALLOWED_EXT

def is_logged_in():
    return "user" in session

def summary_to_points(summary_text: str):
    if not summary_text:
        return []
    lines = [ln.strip() for ln in summary_text.splitlines() if ln.strip()]
    points = []
    for ln in lines:
        ln = re.sub(r"^\s*[\-\*\u2022]\s*", "", ln)
        ln = re.sub(r"^\s*\d+[\.\)]\s*", "", ln)
        if ln:
            points.append(ln)
    return points or [summary_text.strip()]

def redact_phi(text: str) -> str:
    if not text:
        return text

    redacted = text
    patterns = [
        r"(?im)^\s*patient\s*name\s*:\s*.*$",
        r"(?im)^\s*name\s*:\s*.*$",
        r"(?im)^\s*patient\s*id\s*:\s*.*$",
        r"(?im)^\s*(mrn|uhid)\s*:\s*.*$",
        r"(?im)^\s*date\s*:\s*.*$",
        r"(?im)\b\d{1,2}[-/][A-Za-z]{3}[-/]\d{2,4}\b",
        r"(?im)\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b",
        r"(?im)^\s*radiologist\s*:\s*.*$",
        r"(?im)^\s*doctor\s*:\s*.*$",
        r"(?im)^\s*consultant\s*:\s*.*$",
        r"(?im)^\s*signature\s*:\s*.*$",
        r"(?im)\bdr\.\s*[A-Za-z .]+\b",
    ]
    for pat in patterns:
        redacted = re.sub(pat, "[REDACTED]", redacted)

    redacted = re.sub(r"\b[A-Z]{1,5}-\d{2,6}(-\d{1,6})?\b", "[REDACTED_ID]", redacted)
    redacted = re.sub(r"\b(\+91[-\s]?)?[6-9]\d{9}\b", "[REDACTED_PHONE]", redacted)
    redacted = re.sub(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b", "[REDACTED_EMAIL]", redacted)

    return redacted

def _env_or_which(env_key: str, exe_name: str) -> str:
    val = (os.getenv(env_key) or "").strip()
    if val and os.path.exists(val):
        return val
    found = shutil.which(exe_name) or ""
    return found

def _ensure_ghostscript_on_path():
    """
    Some stacks expect 'gs' or 'gswin64c'. If user sets GHOSTSCRIPT_PATH to the bin folder,
    we add it to PATH for the process. If gs is already in PATH, do nothing.
    """
    if shutil.which("gs") or shutil.which("gswin64c"):
        return

    gs_bin = (os.getenv("GHOSTSCRIPT_PATH") or "").strip()  # optional: bin folder or full exe
    if not gs_bin:
        return

    # allow both "...\bin" or "...\bin\gswin64c.exe"
    if gs_bin.lower().endswith(".exe"):
        gs_bin = os.path.dirname(gs_bin)

    if os.path.exists(gs_bin):
        os.environ["PATH"] = gs_bin + os.pathsep + os.environ.get("PATH", "")

def _looks_like_header_only(t: str) -> bool:
    t_low = (t or "").lower()

    keywords = ["impression", "findings", "history", "procedure", "indication", "technique", "conclusion"]
    has_keyword = any(k in t_low for k in keywords)

    header_signals = ["www.", "fax", "suite", "road", "street", "phone", "tel", "completing your image"]
    has_header_signal = any(s in t_low for s in header_signals)

    word_count = len((t or "").split())
    return (word_count < 80) or (has_header_signal and not has_keyword)

def extract_text(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".txt":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(path)
            return "\n".join([p.text for p in doc.paragraphs]).strip()
        except Exception as e:
            print("[DOCX ERROR]", repr(e))
            return ""

    if ext == ".pdf":
        text = ""

        # 1) pypdf (fast for text PDFs)
        try:
            from pypdf import PdfReader
            reader = PdfReader(path)
            parts = [(page.extract_text() or "") for page in reader.pages]
            text = "\n".join(parts).strip()
            print(f"[INFO] pypdf extracted chars: {len(text)}")
        except Exception as e:
            print("[PDF TEXT ERROR - pypdf]", repr(e))
            text = ""

        # Helper: decide if we should OCR
        def looks_like_scanned_or_header_only(t: str) -> bool:
            t_low = (t or "").lower()
            keywords = ["impression", "findings", "history", "procedure", "indication", "technique", "conclusion"]
            has_keyword = any(k in t_low for k in keywords)
            header_signals = ["www.", "fax", "suite", "road", "street", "phone", "tel", "completing your image"]
            has_header_signal = any(s in t_low for s in header_signals)
            word_count = len((t or "").split())
            return (word_count < 80) or (has_header_signal and not has_keyword)

        # 2) OCR using PyMuPDF rendering (NO poppler/ghostscript needed)
        if looks_like_scanned_or_header_only(text):
            print("[INFO] PDF looks scanned/header-only -> using PyMuPDF OCR...")

            try:
                import fitz  # PyMuPDF
                import pytesseract
                from PIL import Image, ImageOps, ImageEnhance
                import io

                # set tesseract from .env if provided
                tess_cmd = (os.getenv("TESSERACT_CMD") or "").strip()
                if tess_cmd and os.path.exists(tess_cmd):
                    pytesseract.pytesseract.tesseract_cmd = tess_cmd
                elif tess_cmd:
                    print("[WARN] TESSERACT_CMD does not exist:", tess_cmd)

                doc = fitz.open(path)

                ocr_parts = []
                # render scale: 2.5 to 3.0 is usually good (higher = clearer but slower)
                zoom = 3.0
                mat = fitz.Matrix(zoom, zoom)

                # OCR configs to try
                psm_candidates = [3, 4, 6, 11]

                for i in range(len(doc)):
                    page = doc.load_page(i)
                    pix = page.get_pixmap(matrix=mat, alpha=False)

                    img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")

                    # preprocessing
                    gray = ImageOps.grayscale(img)
                    gray = ImageEnhance.Contrast(gray).enhance(2.0)
                    gray = ImageEnhance.Sharpness(gray).enhance(1.5)
                    # binarize
                    gray = gray.point(lambda x: 255 if x > 170 else 0, mode="1")

                    page_best = ""
                    page_best_len = 0

                    for psm in psm_candidates:
                        cfg = f"--oem 3 --psm {psm}"
                        out = (pytesseract.image_to_string(gray, lang="eng", config=cfg) or "").strip()
                        if len(out) > page_best_len:
                            page_best = out
                            page_best_len = len(out)

                    print(f"[INFO] OCR page {i+1} best chars: {page_best_len}")
                    if page_best:
                        ocr_parts.append(page_best)

                ocr_text = "\n\n".join(ocr_parts).strip()
                print(f"[INFO] OCR total chars (PyMuPDF): {len(ocr_text)}")

                if len(ocr_text) > len(text):
                    text = ocr_text

            except Exception as e:
                print("[OCR ERROR - PyMuPDF]", repr(e))

        return text

    # -------------------------
    # Images OCR
    # -------------------------
    if ext in {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}:
        try:
            from PIL import Image
            import pytesseract

            tess_cmd = (os.getenv("TESSERACT_CMD") or "").strip()
            if tess_cmd and os.path.exists(tess_cmd):
                pytesseract.pytesseract.tesseract_cmd = tess_cmd

            img = Image.open(path).convert("RGB")
            cfg = r"--oem 3 --psm 6"
            text = (pytesseract.image_to_string(img, lang="eng", config=cfg) or "").strip()
            print(f"[INFO] Image OCR chars: {len(text)}")
            return text
        except Exception as e:
            print("[IMAGE OCR ERROR]", repr(e))
            return ""

    return ""

# ---------------------------------------------------------
# Summarizer PKL parsing
# ---------------------------------------------------------
class SimpleDocument:
    def __init__(self, page_content: str, metadata: dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}

def parse_biolaysumm_text(raw: str):
    if not raw:
        return None, None

    m = re.search(
        r"Radiology report\s*\(clinical\)\s*(.*?)\s*Layman report\s*\(non-clinical\)\s*(.*)$",
        raw,
        flags=re.S | re.I,
    )
    if not m:
        return None, None

    clinical = (m.group(1) or "").strip()
    layman = (m.group(2) or "").strip()
    if not clinical or not layman:
        return None, None
    return clinical, layman

def load_summarization_examples(pkl_path: str):
    if not os.path.exists(pkl_path):
        raise RuntimeError(f"Missing summarization PKL at: {pkl_path}")

    with open(pkl_path, "rb") as f:
        items = pickle.load(f)

    docs = []
    for i, it in enumerate(items):
        raw = (it or {}).get("text") if isinstance(it, dict) else str(it)
        clinical, layman = parse_biolaysumm_text(raw)
        if not clinical or not layman:
            continue
        meta = {"id": i, "clinical": clinical, "layman": layman}
        docs.append(SimpleDocument(page_content=clinical, metadata=meta))

    if not docs:
        raise RuntimeError("No valid clinical/layman pairs found inside the PKL.")
    return docs

# ---------------------------------------------------------
# Init Q&A store (src)
# ---------------------------------------------------------
def init_qa_vectorstore():
    store = QAFaissVectorStore(persist_dir=PERSIST_DIR, embedding_model=EMBEDDING_MODEL)

    qa_index_path = os.path.join(PERSIST_DIR, QA_INDEX_NAME)
    qa_meta_path = os.path.join(PERSIST_DIR, QA_META_NAME)

    if os.path.exists(qa_index_path) and os.path.exists(qa_meta_path):
        store.load()
    else:
        print("[INFO] Building Q&A FAISS store from data/ ...")
        docs = load_qa_documents(os.path.join(BASE_DIR, "data"))
        if not docs:
            raise RuntimeError("No documents found in data/. Add files and restart.")
        store.build_from_documents(docs)

    return store

# ---------------------------------------------------------
# Init Summarizer store (src_code)
# ---------------------------------------------------------
def init_summarizer_vectorstore():
    os.makedirs(SUMMARY_DIR, exist_ok=True)

    store = SummFaissVectorStore(persist_dir=SUMMARY_DIR, embedding_model=EMBEDDING_MODEL)

    summ_index_path = os.path.join(SUMMARY_DIR, SUMM_INDEX_NAME)
    summ_meta_path = os.path.join(SUMMARY_DIR, SUMM_META_NAME)

    if os.path.exists(summ_index_path) and os.path.exists(summ_meta_path):
        store.load()
        return store

    print("[INFO] Building Summarizer FAISS store from PKL inside faiss_store_2 ...")
    docs = load_summarization_examples(SUMMARY_PKL_PATH)
    store.build_from_documents(docs)
    return store

qa_store = init_qa_vectorstore()
summ_store = init_summarizer_vectorstore()

# ---------------------------------------------------------
# Medical QA logic
# ---------------------------------------------------------
MEDICAL_KEYWORDS = {
    "disease","symptom","symptoms","diagnosis","diagnose","treatment","treat","therapy",
    "medicine","medication","drug","dose","dosage","tablet","capsule","injection",
    "side effect","side effects","contraindication","prevention","infection","virus",
    "bacteria","fungal","parasite","fever","cough","pain","headache","vomiting","nausea",
    "diarrhea","fatigue","rash","allergy","diabetes","hypertension","blood pressure","bp",
    "cancer","tumor","asthma","pneumonia","tb","heart","cardiac","stroke","kidney","liver",
    "lungs","brain","skin","pregnancy","period","menstrual","cholesterol","thyroid","anemia",
    "vaccine","vaccination","immunization","xray","x-ray","ct","mri","ultrasound","scan",
    "report","lab","test","cbc","lft","kft","doctor","hospital","clinic","surgery","operation",
    "clinical","medical","patient"
}
MEDICAL_PATTERNS = [
    r"\b(symptom|symptoms|causes|treatment|prevention|risk factors|complications)\b",
    r"\b(report|scan|x[- ]?ray|mri|ct|ultrasound|lab test|blood test)\b",
    r"\b(dose|dosage|side effects?|contraindications?)\b",
]

def is_medical_question(q: str) -> bool:
    q = (q or "").strip().lower()
    if not q:
        return False
    for kw in MEDICAL_KEYWORDS:
        if kw in q:
            return True
    for pat in MEDICAL_PATTERNS:
        if re.search(pat, q):
            return True
    return False

def non_medical_reply(question: str) -> str:
    return (
        f"Q: {question}\n"
        "A: I can answer only medical/health-related questions. "
        "Please ask about diseases, symptoms, medicines, lab tests, reports, treatments, "
        "prevention, or risk factors."
    )

def rag_answer(query: str, top_k: int = TOP_K_DEFAULT) -> str:
    results = qa_store.query(query, top_k=top_k)

    contexts = []
    used = 0
    MAX_CONTEXT_CHARS = 16000

    for r in results:
        meta = r.get("metadata") or {}
        t = (meta.get("text") or "").strip()
        if not t:
            continue
        remaining = MAX_CONTEXT_CHARS - used
        if remaining <= 0:
            break
        piece = t[:remaining]
        used += len(piece)
        contexts.append(piece)

    if not contexts:
        return f"Q: {query}\nA: I don't know based on the available medical data."

    context_block = "\n\n---\n\n".join(contexts)

    prompt = f"""
You are a medical question-answering system.

The context below contains medical Q&A examples from a medical dataset (MedQuad).
Answer the user's question in the SAME style.

STRICT OUTPUT FORMAT:
Q: <repeat the user's question exactly>
A: <write a detailed medical answer in plain language>

Rules:
- Do NOT mention sources.
- Do NOT say "based on the provided sources".
- Do NOT add "Sources used".
- Do NOT invent facts. Use ONLY the context.
- If answer is not present in context:
  Q: <question>
  A: I don't know based on the available medical data.

CONTEXT:
{context_block}

USER QUESTION:
{query}

OUTPUT:
Q: {query}
A:
""".strip()

    resp = llm.invoke(prompt)
    return resp.content.strip()

# ---------------------------------------------------------
# ✅ Summarization logic (Groq ALWAYS used)
# ---------------------------------------------------------
def rag_layman_summary(report_text: str, top_k: int = TOP_K_SUMMARY) -> str:
    report_text = (report_text or "").strip()
    if not report_text:
        return "I couldn't read any text from the uploaded report."

    results = summ_store.query(report_text, top_k=top_k)

    examples = []
    for r in results:
        meta = r.get("metadata") or {}
        clinical = (meta.get("clinical") or "").strip()
        layman = (meta.get("layman") or "").strip()
        if clinical and layman:
            examples.append((clinical, layman))

    if not examples:
        prompt = f"""
You are a helpful assistant that converts radiology reports into simple, non-clinical language.

RULES:
- Use simple English, short sentences.
- Do NOT diagnose.
- Do NOT guess missing findings.
- If the text looks incomplete (only hospital address/phone/website), say: "Report text looks incomplete."
- Output 5-10 bullet points.

REPORT TEXT:
{report_text}

OUTPUT (bullets only):
""".strip()
        resp = llm.invoke(prompt)
        return resp.content.strip()

    ex_block = "\n\n---\n\n".join(
        [f"EXAMPLE {i+1}\nClinical:\n{c}\n\nLayman:\n{l}" for i, (c, l) in enumerate(examples)]
    )

    prompt = f"""
You convert radiology reports into simple, non-clinical language for patients.

TASK:
Given the NEW clinical report, write a layman summary like the examples.

STYLE RULES:
- Use simple English, short sentences.
- Do NOT diagnose. Do NOT guess missing findings.
- Output 5-10 bullet points (each bullet 1 line).
- Only mention follow-up advice if it is clearly stated in the report.

EXAMPLES (clinical -> layman):
{ex_block}

NEW CLINICAL REPORT:
{report_text}

OUTPUT (bullets only):
""".strip()

    resp = llm.invoke(prompt)
    return resp.content.strip()

# ---------------------------------------------------------
# Routes
# ---------------------------------------------------------
@app.route("/", methods=["GET"])
def home():
    if not is_logged_in():
        return redirect(url_for("signin"))
    return render_template("index.html", user=session.get("user"))

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if is_logged_in():
        return redirect(url_for("home"))

    if request.method == "POST":
        name = (request.form.get("name") or "").strip()
        email = (request.form.get("email") or "").strip().lower()
        password = (request.form.get("password") or "").strip()

        if not name or not email or not password:
            flash("All fields are required.", "error")
            return render_template("signup.html", user=session.get("user"))

        try:
            supabase.auth.sign_up({
                "email": email,
                "password": password,
                "options": {"data": {"name": name}}
            })
            flash("Account created! Please sign in (verify email if required).", "success")
            return redirect(url_for("signin"))
        except Exception as e:
            flash(f"Signup failed: {str(e)}", "error")

    return render_template("signup.html", user=session.get("user"))

@app.route("/signin", methods=["GET", "POST"])
def signin():
    if is_logged_in():
        return redirect(url_for("home"))

    if request.method == "POST":
        email = (request.form.get("email") or "").strip().lower()
        password = (request.form.get("password") or "").strip()

        if not email or not password:
            flash("Email and password are required.", "error")
            return render_template("signin.html", user=session.get("user"))

        try:
            res = supabase.auth.sign_in_with_password({"email": email, "password": password})
            user = res.user
            sess = res.session

            session["access_token"] = sess.access_token
            session["refresh_token"] = sess.refresh_token
            session["user"] = {
                "id": user.id,
                "email": user.email,
                "name": (user.user_metadata or {}).get("name", "User"),
            }

            flash("Signed in successfully.", "success")
            return redirect(url_for("home"))

        except Exception as e:
            flash(f"Invalid email/password (or email not verified). {str(e)}", "error")

    return render_template("signin.html", user=session.get("user"))

@app.route("/logout", methods=["GET"])
def logout():
    session.clear()
    flash("Logged out successfully.", "success")
    return redirect(url_for("signin"))

@app.route("/report-summarizer", methods=["GET", "POST"])
def report_summarizer_page():
    if not is_logged_in():
        return redirect(url_for("signin"))

    if request.method == "POST":
        f = request.files.get("file") or request.files.get("report_file")
        if not f or f.filename == "":
            return render_template("report_summarizer.html", user=session.get("user"), error="No file selected.")

        if not allowed_file(f.filename):
            return render_template(
                "report_summarizer.html",
                user=session.get("user"),
                error=f"Unsupported file type. Allowed: {', '.join(sorted(ALLOWED_EXT))}"
            )

        filename = secure_filename(f.filename)
        save_path = os.path.join(UPLOAD_DIR, filename)
        f.save(save_path)

        extracted_text = extract_text(save_path)
        safe_text = redact_phi(extracted_text)

        summary_text = rag_layman_summary(safe_text, top_k=TOP_K_SUMMARY)
        summary_points = summary_to_points(summary_text)

        return render_template(
            "report_summarizer.html",
            user=session.get("user"),
            extracted_text=safe_text,
            summary=summary_points,
            filename=filename,
        )

    return render_template("report_summarizer.html", user=session.get("user"))

@app.route("/medical-qa", methods=["GET", "POST"], endpoint="medical_qa_chat")
def medical_qa_page():
    if not is_logged_in():
        return redirect(url_for("signin"))

    if "chat_history" not in session:
        session["chat_history"] = []

    if request.method == "POST" and request.form.get("reset_chat") == "1":
        session["chat_history"] = []
        return redirect(url_for("medical_qa_chat"))

    if request.method == "POST":
        question = (request.form.get("question") or "").strip()
        if question:
            history = session["chat_history"]
            history.append({"role": "user", "content": question})

            if not is_medical_question(question):
                answer = non_medical_reply(question)
            else:
                answer = rag_answer(question, top_k=TOP_K_DEFAULT)

            history.append({"role": "assistant", "content": answer})
            session["chat_history"] = history

        return redirect(url_for("medical_qa_chat"))

    return render_template(
        "medical_qa.html",
        user=session.get("user"),
        chat_history=session.get("chat_history", [])
    )

@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})

if __name__ == "__main__":
    app.run()
